from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from period import Period


WEEKDAYS = [
    'Понедельник',
    'Вторник',
    'Среда',
    'Четверг',
    'Пятница',
    'Суббота'
]

PERIODS = [
    '  9:30 - 11:00',
    '11:10 - 12:40',
    '13:00 - 14:30',
    '15:00 - 16:30',
    '16:40 - 18:10',
    '18:30 - 20:00',
    '20:10 - 21:40'
]


def head_cell_format(cell, string):
    cell.autofit = True
    cell.width = Cm(1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(string)
    font = p.runs[0].font
    font.bold = True
    font.size = Pt(12)


def merge_and_fill_period_cells(cells):
    period_cells = [cells[i].merge(cells[i + 1])
                    for i in range(1, len(cells)-1, 2)]
    for i in range(len(period_cells)):
        cell = period_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        head_cell_format(cell, str(i + 1))
        cell.add_paragraph(PERIODS[i])
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[1].runs[0].font.size = Pt(10)
        cell.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def get_periods(html):
    # Список дней
    days = html.find_all('h3')
    if days[0].text not in WEEKDAYS:
        days = days[1:]

    # Список пар
    std = {i: [] for i in range(1, 7)}
    for day in days:
        study_day = day
        while study_day.next_sibling and study_day.next_sibling not in days:
            study_day = study_day.next_sibling
            if ' пара (' in study_day.text:
                std[WEEKDAYS.index(day.text) + 1].append(study_day)

    periods = []
    for day in std:
        period = std.get(day)
        if period:
            for lesson in period:
                periods.append(Period(lesson, period, day, WEEKDAYS))

    return periods


# Заполнение таблицы
def filling(table, html):
    def filling_into_cell(cell, period, ind=0):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        name = period.subjects[ind].name
        subject_type = period.subjects[ind].type
        groups = period.subjects[ind].groups
        groups_fmt = groups[0] if len (groups) <= 1 else f'{groups[0]}, ...'
        lector = period.subjects[ind].lector

        cell.paragraphs[0].add_run(name)
        rc = cell.paragraphs[0].runs[0]
        rc.font.bold = True
        rc.font.size = Pt(12)

        cell.paragraphs[0].add_run(' ' + subject_type)
        rc = cell.paragraphs[0].runs[1]
        rc.font.bold = False
        rc.font.size = Pt(8)

        cell.add_paragraph(groups_fmt)
        rc = cell.paragraphs[1].runs[0]
        rc.font.bold = True
        rc.font.size = Pt(12)

        if lector != None:
            cell.add_paragraph(lector)
            rc = cell.paragraphs[2].runs[0]
            rc.font.bold = False
            rc.font.size = Pt(10)

        for cell in cell.paragraphs:
            cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraph_format.line_spacing = .8

    periods = get_periods(html)

    for period in periods:
        column = table.columns[period.day]
        cell_number = period.number + period.number - 1
        cell = column.cells[cell_number]

        if period.count == 2:
            filling_into_cell(cell, period)
            filling_into_cell(column.cells[cell_number + 1], period, 1)
        elif period.subjects[0].turn == 0:
            cell.merge(column.cells[cell_number + 1])
            filling_into_cell(cell, period)
        elif period.subjects[0].turn == 1:
            filling_into_cell(cell, period)
        else:
            filling_into_cell(column.cells[cell_number + 1], period)


# Объект таблицы
class ScheduleTable:
    def __init__(self, document, html):
        self.table = document.add_table(15, 7)
        self.table.style = 'Table Grid'
        self.table.alignment = WD_TABLE_ALIGNMENT.CENTER

        head_cells = self.table.rows[0].cells
        for i, item in enumerate(['Пара', *WEEKDAYS]):
            head_cell_format(head_cells[i], item)

        merge_and_fill_period_cells(self.table.columns[0].cells)
        for row in self.table.rows[1:]:
            row.height = Cm(1.4)

        filling(self.table, html)
